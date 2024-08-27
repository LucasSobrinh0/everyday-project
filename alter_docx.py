from docx import Document

# Lista com os nomes dos arquivos
arquivos = ["um.docx", "dois.docx", "tres.docx"]

# Definir a frase a ser substituída e a nova frase
frase_antiga = "Agora sim. Essa é a terceira linha do documento =D"
frase_nova = "Essa é a terceira linha do documento. Que bom que você está lendo até agora!!!!"

# Palavra que deve ficar em negrito
palavra_em_negrito = "terceira"

# Iterar sobre os arquivos e fazer a substituição
for arquivo in arquivos:
    doc = Document(arquivo)
    for para in doc.paragraphs:
        if frase_antiga in para.text:
            # Substituir a frase antiga pela nova
            para.text = para.text.replace(frase_antiga, frase_nova)
            
            # Encontrar e formatar a palavra em negrito
            if palavra_em_negrito in para.text:
                partes = para.text.split(palavra_em_negrito)
                para.clear()  # Limpar o texto do parágrafo
                
                # Adicionar as partes de volta, com a palavra em negrito
                for i, parte in enumerate(partes):
                    para.add_run(parte)
                    if i < len(partes) - 1:
                        run = para.add_run(palavra_em_negrito)
                        run.bold = True

    doc.save(f"atualizado_{arquivo}")
